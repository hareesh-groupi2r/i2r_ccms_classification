"""
Enhanced Document Type Classification Service
Supports automatic detection of 10+ document types using keyword and pattern matching
"""

import re
import sys
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from pypdf import PdfReader
from docx import Document

from .interfaces import IDocumentTypeService, ProcessingResult, ProcessingStatus, DocumentType
from .configuration_service import get_config_service

# Add classifier directory to path
import os
classifier_path = os.path.join(os.path.dirname(__file__), '..', '..', 'classifier')
if classifier_path not in sys.path:
    sys.path.insert(0, classifier_path)

from pdf_extractor import PDFExtractor


class DocumentTypeService(IDocumentTypeService):
    """Enhanced service for classifying multiple document types"""
    
    def __init__(self, config_service=None):
        self.config_service = config_service or get_config_service()
        self.config = self.config_service.get_service_config("document_type")
        
        # Extract configuration for all document types
        self.document_types_config = self._load_document_type_configs()
        self.pages_to_check = self.config.get("pages_to_check", 1)  # Only first page for fast processing
        self.confidence_threshold = self.config.get("confidence_threshold", 0.6)
        self.min_keyword_matches = self.config.get("min_keyword_matches", 2)
        self.pattern_weight = self.config.get("pattern_weight", 3)
        self.keyword_weight = self.config.get("keyword_weight", 1)
        self.filename_weight = self.config.get("filename_weight", 2)
        self.length_based_scoring = self.config.get("length_based_scoring", True)
        self.document_type_priority = self.config.get("document_type_priority", {})
        self.filename_keywords = self.config.get("filename_keywords", {})
        
        # Initialize PDFExtractor with OCR support for scanned documents
        self.pdf_extractor = PDFExtractor(max_pages=self.pages_to_check, ocr_threshold=50)
    
    def _load_document_type_configs(self) -> Dict[DocumentType, Dict[str, any]]:
        """Load configuration for all document types"""
        document_configs = {}
        
        # Map DocumentType enum to config keys
        type_mappings = {
            DocumentType.CORRESPONDENCE: "correspondence",
            DocumentType.MEETING_MINUTES: "meeting_minutes", 
            DocumentType.PROGRESS_REPORTS: "progress_reports",
            DocumentType.CHANGE_ORDERS: "change_orders",
            DocumentType.CONTRACT_AGREEMENTS: "contract_agreements",
            DocumentType.PAYMENT_STATEMENTS: "payment_statements",
            DocumentType.COURT_ORDERS: "court_orders",
            DocumentType.POLICY_CIRCULARS: "policy_circulars",
            DocumentType.TECHNICAL_DRAWINGS: "technical_drawings"
        }
        
        for doc_type, config_key in type_mappings.items():
            keywords = self.config.get(f"{config_key}_keywords", [])
            patterns = self.config.get(f"{config_key}_patterns", {})
            
            document_configs[doc_type] = {
                "keywords": keywords,
                "patterns": self._parse_search_terms(patterns)
            }
        
        return document_configs
    
    def _parse_search_terms(self, search_terms_config: Dict[str, str]) -> Dict[str, int]:
        """Parse search terms configuration with regex flags"""
        parsed_terms = {}
        for pattern, flags_str in search_terms_config.items():
            if flags_str == "MULTILINE":
                parsed_terms[pattern] = re.MULTILINE
            else:
                parsed_terms[pattern] = 0
        return parsed_terms
    
    def _get_pdf_page_count(self, file_path: str) -> int:
        """Get the total number of pages in a PDF document"""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                return len(reader.pages)
        except Exception as e:
            print(f"Warning: Could not get page count for {file_path}: {e}")
            return 0
    
    def _calculate_filename_score(self, file_path: str) -> Dict[DocumentType, int]:
        """Calculate scores based on filename keywords"""
        filename = Path(file_path).stem.lower()  # Get filename without extension
        filename_scores = {}
        
        # Initialize all document types with 0 score
        for doc_type in DocumentType:
            if doc_type not in [DocumentType.UNKNOWN, DocumentType.OTHERS]:
                filename_scores[doc_type] = 0
        
        # Check filename against keywords for each document type
        for doc_type_key, keywords in self.filename_keywords.items():
            try:
                doc_type = DocumentType(doc_type_key)
                score = 0
                for keyword in keywords:
                    if keyword.lower() in filename:
                        score += 1
                filename_scores[doc_type] = score
            except ValueError:
                continue  # Skip invalid document types
        
        return filename_scores

    def classify_document(self, file_path: str, **kwargs) -> ProcessingResult:
        """
        Classify document type from PDF or DOCX file
        
        Args:
            file_path: Path to the PDF or DOCX file
            **kwargs: Additional options
                - pages_to_check: Number of pages to analyze
                - use_advanced_classification: Use regex patterns for enhanced classification
        
        Returns:
            ProcessingResult with DocumentType and confidence score
        """
        try:
            # Validate file exists
            if not Path(file_path).exists():
                return ProcessingResult(
                    status=ProcessingStatus.ERROR,
                    error_message=f"File not found: {file_path}"
                )
            
            # Extract text based on file type
            pages_to_check = kwargs.get("pages_to_check", self.pages_to_check)
            
            # Get filename scores for enhanced classification
            filename_scores = self._calculate_filename_score(file_path)
            
            # Always use 1 page for fast document type classification
            # Dynamic page limits are only used during actual document processing
            
            # Determine file type and extract text accordingly
            file_extension = Path(file_path).suffix.lower()
            page_count = 0
            
            if file_extension == '.docx':
                text_result = self._extract_text_from_docx(file_path)
            else:
                # Get PDF page count for contract agreement detection
                page_count = self._get_pdf_page_count(file_path)
                text_result = self._extract_text_from_pdf(file_path, pages_to_check)
            
            if text_result.status == ProcessingStatus.ERROR:
                return text_result
            
            text_content = text_result.data
            
            # Classify based on text content with page count and filename scores
            classification_result = self.classify_from_text(
                text_content, 
                page_count=page_count,
                filename_scores=filename_scores,
                use_advanced_classification=kwargs.get("use_advanced_classification", True)
            )
            
            # Add file metadata to result
            classification_result.metadata = classification_result.metadata or {}
            classification_result.metadata.update({
                "file_path": file_path,
                "pages_analyzed": pages_to_check,
                "text_length": len(text_content),
                "total_pages": page_count
            })
            
            return classification_result
            
        except Exception as e:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                error_message=f"Error classifying document: {str(e)}"
            )
    
    def classify_from_text(self, text_content: str, page_count: int = 0, **kwargs) -> ProcessingResult:
        """
        Classify document type from text content using enhanced multi-type classification
        
        Args:
            text_content: Text content to analyze
            **kwargs: Additional options
                - use_advanced_classification: Use regex patterns for detailed analysis
                - confidence_threshold: Minimum confidence for classification
        
        Returns:
            ProcessingResult with DocumentType and confidence score
        """
        try:
            if not text_content or not text_content.strip():
                return ProcessingResult(
                    status=ProcessingStatus.ERROR,
                    error_message="No text content provided for classification"
                )
            
            text_lower = text_content.lower()
            use_advanced = kwargs.get("use_advanced_classification", True)
            confidence_threshold = kwargs.get("confidence_threshold", self.confidence_threshold)
            
            # Get filename scores if provided
            filename_scores = kwargs.get("filename_scores", {})
            
            # Calculate scores for all document types
            type_scores = {}
            detailed_analysis = {}
            
            for doc_type, config in self.document_types_config.items():
                keyword_score = self._calculate_keyword_score(text_lower, config["keywords"])
                pattern_score = 0
                filename_score = filename_scores.get(doc_type, 0)
                
                if use_advanced and config["patterns"]:
                    pattern_score = self._calculate_pattern_score(text_content, config["patterns"])
                
                # Combined score with weights
                total_score = (keyword_score * self.keyword_weight) + (pattern_score * self.pattern_weight) + (filename_score * self.filename_weight)
                
                # Apply length-based scoring for specific document types
                if self.length_based_scoring:
                    length_bonus = self._calculate_length_bonus(doc_type, text_content, page_count)
                    total_score += length_bonus
                
                type_scores[doc_type] = total_score
                detailed_analysis[doc_type.value] = {
                    "keyword_score": keyword_score,
                    "pattern_score": pattern_score,
                    "filename_score": filename_score,
                    "total_score": total_score,
                    "keyword_matches": self._get_keyword_matches(text_lower, config["keywords"]),
                    "pattern_matches": self._get_pattern_matches(text_content, config["patterns"]) if use_advanced else []
                }
            
            # Determine best match
            if not any(score > 0 for score in type_scores.values()):
                # No matches found - return OTHERS
                doc_type = DocumentType.OTHERS
                confidence = 0.4  # Low confidence fallback
            else:
                # Find document type with highest score
                max_score = max(type_scores.values())
                best_matches = [doc_type for doc_type, score in type_scores.items() if score == max_score]
                
                if len(best_matches) == 1:
                    doc_type = best_matches[0]
                else:
                    # Handle ties using priority system
                    doc_type = self._resolve_tie(best_matches)
                
                # Calculate confidence based on score and separation from next best
                scores_list = sorted(type_scores.values(), reverse=True)
                confidence = self._calculate_confidence(max_score, scores_list)
            
            # Apply minimum threshold check
            if confidence < confidence_threshold:
                doc_type = DocumentType.UNKNOWN
                confidence = confidence  # Keep original confidence for analysis
            
            # Check for unsupported document types and add warnings
            warning_message = self._get_support_warning(doc_type)
            
            return ProcessingResult(
                status=ProcessingStatus.SUCCESS,
                data=doc_type,
                confidence=confidence,
                metadata={
                    "all_scores": {dt.value: score for dt, score in type_scores.items()},
                    "detailed_analysis": detailed_analysis,
                    "text_length": len(text_content),
                    "total_pages": page_count,
                    "advanced_patterns_used": use_advanced,
                    "confidence_threshold": confidence_threshold,
                    "classification_method": "enhanced_multi_type",
                    "support_warning": warning_message
                }
            )
            
        except Exception as e:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                error_message=f"Error in text classification: {str(e)}"
            )
    
    def _calculate_keyword_score(self, text_lower: str, keywords: List[str]) -> int:
        """Calculate score based on keyword matches"""
        score = 0
        for keyword in keywords:
            if keyword.lower() in text_lower:
                score += 1
        return score
    
    def _calculate_pattern_score(self, text: str, search_terms: Dict[str, int]) -> int:
        """Calculate score based on regex pattern matches"""
        score = 0
        for pattern, flags in search_terms.items():
            try:
                if re.search(pattern, text, re.IGNORECASE | flags):
                    score += 1
            except re.error as e:
                print(f"Warning: Invalid regex pattern '{pattern}': {e}")
                continue
        return score
    
    def _calculate_length_bonus(self, doc_type: DocumentType, text_content: str, page_count: int = 0) -> float:
        """Apply length-based scoring bonuses for specific document types"""
        text_length = len(text_content)
        
        # Contract agreements: Page count is the primary indicator
        if doc_type == DocumentType.CONTRACT_AGREEMENTS:
            if page_count > 200:  # Strong indicator for contract agreements (200+ pages)
                return 3.0  # Very high bonus for 200+ page documents
            elif page_count > 100:  # Large document but may not be contract
                return 1.0
            elif text_length > 50000:  # Very large text content
                return 1.5
            elif text_length > 20000:  # Medium bonus for large docs
                return 0.5
            else:
                return -1.0  # Penalty for short documents - likely not contracts
        
        # Correspondence letters: Favor typical letter lengths
        elif doc_type == DocumentType.CORRESPONDENCE:
            if 500 <= text_length <= 8000:  # Expanded typical letter range
                return 1.0  # Increased bonus for correspondence
            elif text_length > 15000:  # Too long for typical letter
                return -0.3  # Light penalty for very long letters
            else:
                return 0.2  # Small bonus for short notes/memos
        
        # Meeting minutes: Structured content, medium length
        elif doc_type == DocumentType.MEETING_MINUTES:
            if 2000 <= text_length <= 15000:  # Expanded range for minutes
                return 0.8
            elif text_length < 1000:  # Too short for proper minutes
                return -0.5
        
        # Technical drawings often have shorter text with specific terms
        elif doc_type == DocumentType.TECHNICAL_DRAWINGS and text_length < 2000:
            return 0.5
        
        # Progress reports: Medium to long structured documents
        elif doc_type == DocumentType.PROGRESS_REPORTS and 3000 <= text_length <= 12000:
            return 0.6
        
        return 0
    
    def _get_support_warning(self, doc_type: DocumentType) -> Optional[str]:
        """Get warning message for unsupported document types"""
        supported_types = {
            DocumentType.CORRESPONDENCE,
            DocumentType.MEETING_MINUTES,
            DocumentType.PROGRESS_REPORTS,
            DocumentType.CONTRACT_AGREEMENTS
        }
        
        if doc_type not in supported_types:
            type_name = doc_type.value.replace('_', ' ').title()
            return (f"⚠️ WARNING: {type_name} document type detected but processing handler "
                   f"is not implemented yet. This document type will be supported in the next phase. "
                   f"Currently supported types: Contract Agreements, Correspondence Letters, "
                   f"Meeting Minutes, and Progress Reports.")
        
        return None
    
    def _resolve_tie(self, tied_types: List[DocumentType]) -> DocumentType:
        """Resolve ties between document types using priority system"""
        priority_scores = {}
        
        for doc_type in tied_types:
            key = doc_type.value.replace("_", " ").replace(" ", "_")
            priority_scores[doc_type] = self.document_type_priority.get(key, 0)
        
        # Return type with highest priority
        return max(priority_scores.keys(), key=lambda x: priority_scores[x])
    
    def _calculate_confidence(self, max_score: int, scores_list: List[int]) -> float:
        """Calculate confidence based on score and separation from alternatives"""
        if max_score == 0:
            return 0.3  # Low confidence for no matches
        
        # Base confidence from score magnitude
        base_confidence = min(0.9, 0.4 + (max_score * 0.1))
        
        # Bonus for clear separation from next best
        if len(scores_list) > 1 and scores_list[0] > scores_list[1]:
            separation = scores_list[0] - scores_list[1]
            separation_bonus = min(0.2, separation * 0.05)
            base_confidence += separation_bonus
        
        return min(0.95, base_confidence)
    
    def _get_keyword_matches(self, text_lower: str, keywords: List[str]) -> List[str]:
        """Get list of matching keywords for analysis"""
        matches = []
        for keyword in keywords:
            if keyword.lower() in text_lower:
                matches.append(keyword)
        return matches
    
    def _get_pattern_matches(self, text: str, search_terms: Dict[str, int]) -> List[str]:
        """Get list of matching patterns for analysis"""
        matches = []
        for pattern, flags in search_terms.items():
            try:
                if re.search(pattern, text, re.IGNORECASE | flags):
                    matches.append(pattern)
            except re.error:
                continue
        return matches
    
    def _extract_text_from_pdf(self, file_path: str, pages_to_check: int) -> ProcessingResult:
        """Extract text from first few pages of PDF using OCR if needed"""
        try:
            # Use PDFExtractor with OCR support for scanned documents
            text_content, extraction_method = self.pdf_extractor.extract_text(file_path)
            
            if not text_content.strip():
                return ProcessingResult(
                    status=ProcessingStatus.ERROR,
                    error_message="No readable text found in PDF"
                )
            
            # Get PDF page count for metadata
            try:
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
            except:
                total_pages = pages_to_check
            
            return ProcessingResult(
                status=ProcessingStatus.SUCCESS,
                data=text_content,
                metadata={
                    "extraction_method": extraction_method,
                    "total_pages": total_pages,
                    "pages_analyzed": pages_to_check,
                    "text_length": len(text_content)
                }
            )
            
        except Exception as e:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                error_message=f"Error reading PDF: {str(e)}"
            )
    def _extract_text_from_docx(self, file_path: str) -> ProcessingResult:
        """Extract text from DOCX document"""
        try:
            doc = Document(file_path)
            text_content = ""
            paragraphs_processed = 0

            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text.strip() + "\n"
                    paragraphs_processed += 1

            # Extract text from tables if any
            tables_processed = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text.strip() + " "
                tables_processed += 1

            if not text_content.strip():
                return ProcessingResult(
                    status=ProcessingStatus.ERROR,
                    error_message="No readable text found in DOCX"
                )

            return ProcessingResult(
                status=ProcessingStatus.SUCCESS,
                data=text_content.strip(),
                confidence=0.95,
                metadata={
                    "paragraphs_processed": paragraphs_processed,
                    "tables_processed": tables_processed,
                    "text_length": len(text_content.strip())
                }
            )

        except Exception as e:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                error_message=f"Error reading DOCX: {str(e)}"
            )

    def get_classification_details(self, file_path: str) -> ProcessingResult:
        """Get detailed classification information for debugging and analysis"""
        result = self.classify_document(file_path, use_advanced_classification=True)
        
        if result.status == ProcessingStatus.SUCCESS:
            # Add more detailed analysis for debugging
            text_result = self._extract_text_from_pdf(file_path, self.pages_to_check)
            if text_result.status == ProcessingStatus.SUCCESS:
                text_content = text_result.data
                
                # Analyze all document types
                enhanced_analysis = {}
                for doc_type, config in self.document_types_config.items():
                    keyword_matches = self._get_keyword_matches(text_content.lower(), config["keywords"])
                    pattern_matches = self._get_pattern_matches(text_content, config["patterns"])
                    
                    enhanced_analysis[doc_type.value] = {
                        "keyword_matches": keyword_matches,
                        "pattern_matches": pattern_matches,
                        "keyword_count": len(keyword_matches),
                        "pattern_count": len(pattern_matches)
                    }
                
                result.metadata["enhanced_analysis"] = enhanced_analysis
                result.metadata["first_200_chars"] = text_content[:200].replace('\n', ' ')
                result.metadata["document_length_category"] = self._categorize_document_length(len(text_content))
        
        return result
    
    def _categorize_document_length(self, length: int) -> str:
        """Categorize document length for analysis"""
        if length < 1000:
            return "short"
        elif length < 3000:
            return "medium"
        elif length < 8000:
            return "long"
        else:
            return "very_long"
    
    def update_classification_rules(self, new_rules: Dict[str, any]) -> ProcessingResult:
        """Update classification rules dynamically"""
        try:
            config_updated = False
            
            # Update keywords for any document type
            for doc_type in DocumentType:
                if doc_type in [DocumentType.UNKNOWN, DocumentType.OTHERS]:
                    continue
                    
                config_key = doc_type.value
                keywords_key = f"{config_key}_keywords"
                patterns_key = f"{config_key}_patterns"
                
                if keywords_key in new_rules:
                    if config_key in self.document_types_config:
                        self.document_types_config[doc_type]["keywords"] = new_rules[keywords_key]
                    config_updated = True
                
                if patterns_key in new_rules:
                    if config_key in self.document_types_config:
                        self.document_types_config[doc_type]["patterns"] = self._parse_search_terms(new_rules[patterns_key])
                    config_updated = True
            
            # Update general settings
            if "confidence_threshold" in new_rules:
                self.confidence_threshold = new_rules["confidence_threshold"]
                config_updated = True
            
            if "pattern_weight" in new_rules:
                self.pattern_weight = new_rules["pattern_weight"]
                config_updated = True
            
            if "keyword_weight" in new_rules:
                self.keyword_weight = new_rules["keyword_weight"]  
                config_updated = True
            
            # Update configuration service if changes were made
            if config_updated:
                self.config_service.update_service_config("document_type", new_rules)
            
            return ProcessingResult(
                status=ProcessingStatus.SUCCESS,
                data=f"Classification rules updated successfully. Updated: {config_updated}"
            )
            
        except Exception as e:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                error_message=f"Error updating classification rules: {str(e)}"
            )
    
    def get_supported_document_types(self) -> List[str]:
        """Get list of supported document types"""
        return [doc_type.value for doc_type in DocumentType if doc_type != DocumentType.UNKNOWN]
    
    def test_classification_accuracy(self, test_files: Dict[str, DocumentType]) -> ProcessingResult:
        """Test classification accuracy against known document types"""
        try:
            results = {
                "total": len(test_files),
                "correct": 0,
                "incorrect": 0,
                "unknown": 0,
                "details": []
            }
            
            for file_path, expected_type in test_files.items():
                classification_result = self.classify_document(file_path)
                
                if classification_result.status == ProcessingStatus.SUCCESS:
                    predicted_type = classification_result.data
                    
                    if predicted_type == expected_type:
                        results["correct"] += 1
                        status = "correct"
                    elif predicted_type == DocumentType.UNKNOWN:
                        results["unknown"] += 1
                        status = "unknown"
                    else:
                        results["incorrect"] += 1
                        status = "incorrect"
                    
                    results["details"].append({
                        "file": file_path,
                        "expected": expected_type.value,
                        "predicted": predicted_type.value,
                        "confidence": classification_result.confidence,
                        "status": status
                    })
                else:
                    results["details"].append({
                        "file": file_path,
                        "expected": expected_type.value,
                        "predicted": "error",
                        "error": classification_result.error_message,
                        "status": "error"
                    })
            
            accuracy = results["correct"] / results["total"] if results["total"] > 0 else 0
            results["accuracy"] = accuracy
            
            return ProcessingResult(
                status=ProcessingStatus.SUCCESS,
                data=results,
                confidence=accuracy
            )
            
        except Exception as e:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                error_message=f"Error testing classification accuracy: {str(e)}"
            )